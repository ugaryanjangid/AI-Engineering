import os
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel

load_dotenv()

my_api_key = os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("API key kaha hai bhai")

client = Groq(api_key=my_api_key)
model = "openai/gpt-oss-120b"


job_description=""""
Description
This role focuses on managing business data, building reports, and developing dashboards that support better decision-making. The candidate will use analytics tools, cloud platforms, and AI technologies to provide valuable insights, reporting solutions, and data-driven suggestions for various business areas.

Key Responsibilities:
• Examine and interpret structured and unstructured business data
• Design dashboards and reports with Power BI and other BI platforms
• Produce visual reports that show trends, patterns, and business performance
• Work with SQL databases to extract, validate, and report on data
• Build and manage ETL workflows for efficient data flow and transformation
• Apply Python for automation, reporting, and analysis tasks
• Collaborate closely with business and technical teams to capture reporting needs
• Provide accurate insights and analytical reports to support business planning
• Develop data models and reporting frameworks to improve analytics performance
• Assist management teams with data-backed recommendations and business insights
• Enhance reporting systems and dashboard usability for end users
• Use AI-powered tools and Generative AI technologies to improve analytics work
• Work on AWS, Azure, and GCP cloud platforms for analytics-related activities
• Spot business trends and opportunities through data review
• Maintain data quality, consistency, and reporting accuracy across projects
• Keep documentation for reports, dashboards, and analytics processes
• Coordinate with different departments on reporting improvements and data support
• Share regular findings and updates with stakeholders and management

Required Skills:
• Solid hands-on knowledge of Power BI, SQL, Python, and dashboard development
• Experience in ETL design and data transformation workflows
• Strong grasp of business intelligence and data reporting principles
• Familiarity with Tableau and other data visualization tools
• Knowledge of data analytics, business reporting, and data modeling methods
• Experience working on cloud platforms such as AWS, Azure, or GCP
• Understanding of AI tools and Generative AI technologies
• Knowledge of predictive analytics and business forecasting approaches
• Strong analytical thinking and problem-solving skills
• Ability to handle large datasets and create meaningful reports
• Good communication and stakeholder management skills
• Ability to explain reports and insights in clear business terms
• Experience in team-focused and deadline-driven settings
• Good documentation and presentation abilities
• Self-driven mindset with the ability to manage tasks independently

Preferred Skills:
• Background in analytics, consulting, or enterprise reporting projects
• Knowledge of advanced dashboard design and reporting practices
• Familiarity with automation and AI-assisted analytics tools
• Exposure to cloud-based analytics and modern BI platforms
• Understanding of KPI monitoring and operational reporting systems
• Experience creating executive dashboards and business presentations
• Knowledge of dashboard tuning and reporting performance improvements
• Exposure to current business analytics technologies and platforms

Education:
B.Tech / BCA / MCA / B.Sc. in Computer Science, IT, or Equivalent Qualification.
"""

class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirments: list[str]
    responsibilities: list[str]

jobd_schema = JobD.model_json_schema()   #jobD ka schema bana liya

system_prompt = f"""
You are an expert HR assistant.

Your job is to analyze job description and extract structured information from them.

Return ONLY valid JSON matching this schema:

{jobd_schema}
IMPORTANT:
Do NOT return the schema itslef.
Do NOT return fields like "properties","title" or "type".
Fill the schema with actual information extracted from the job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information.
"""

user_prompt = f""""
Analyze the following job description:

{job_description}
"""

message_system={
    "role": "system",
    "content": system_prompt
}

message_user={
    "role":"user",
    "content":user_prompt
}

response_format={
    "type":"json_object"
}

messages = [message_system,message_user]

response=client.chat.completions.create(model=model,messages=messages,response_format=response_format)

answer=response.choices[0].message.content

raw_json=answer
# print(raw_json)   o/p in raw json noe we have to convert this into readaable format.


import json
job_data=json.loads(raw_json)

job = JobD(**job_data)   # now it convert in the job_schema format like line no 73

print(job.minimum_experience)
print(job.education_requirments)




#parser_real
class MatchResult(BaseModel):
    score: float
    details: dict

class Experinece(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []

class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None

    total_experience_years: float | None = None

    skills: list[str] = []
    experinece: list[Experinece] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []

resume_schema = Resume.model_json_schema()   # resume schema ban gya using resume class
def final_score(job,resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """
    message={
        "role": "user",
        "content" : prompt
    }
    messages=[message]
    response_format={
        "type": "json_object"
    }
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)
def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume


from pypdf import PdfReader
from docx import Document

def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"


    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"

    return text

def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None




resume_folder = Path("resumes")
all_results=[]
for file_path in resume_folder.iterdir():
    
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue
    print("\nProcessing:", file_path.name)
    resume_text = read_resume(file_path)
    parsed_resume=parse_resume(resume_text) # llm call1
    time.sleep(5)
    result = final_score(job, parsed_resume) #llm caLL2
    #score and details
    #acount chtgpt
    # request bhejna shhur krega millions
    #chattgot server jam ho jayega
    time.sleep(5)
    print("Score:", result.score)
    all_results.append({
        "name": parsed_resume.name,
        "score": result.score,
        "details": result.details
    })
all_results.sort(
    key=lambda candidate: candidate["score"],
    reverse=True
)
top_2 = all_results[:2]
worst_2 = all_results[-2:]


print("TOP 2 CANDIDATES")
for candidate in top_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )

    print(candidate["details"])

print("LOWEST 2 CANDIDATES")
for candidate in worst_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )
    print(candidate["details"])

    






